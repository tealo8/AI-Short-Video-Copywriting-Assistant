# -*- coding: utf-8 -*-
"""模块6：一键标准化 Word 导出（python-docx）。

规格：A4 版式、微软雅黑全中文统一设置、标题自动编号、
分镜脚本表格化、页脚页码，交付级排版，无格式错乱。
"""
from __future__ import annotations

from io import BytesIO
from typing import Any

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

_PRIMARY = RGBColor(0x1D, 0x1D, 0x1F)
_GRAY = RGBColor(0x86, 0x86, 0x8B)
_BLUE = RGBColor(0x00, 0x71, 0xE3)
_FONT = "微软雅黑"

_TIER_NAMES = {"hot": "热门泛标签", "mid": "行业中标签", "long": "精准长尾标签"}


def _set_font(run, size: float, bold: bool = False, color: RGBColor = _PRIMARY):
    run.font.name = _FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), _FONT)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color


def _heading(doc: Document, text: str, level: int = 1):
    sizes = {0: 16, 1: 13, 2: 11.5}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    _set_font(run, sizes.get(level, 11.5), bold=True, color=_BLUE if level <= 1 else _PRIMARY)
    return p


def _body(doc: Document, text: str, size: float = 10.5, color: RGBColor = _PRIMARY, indent: bool = True):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.6
    if indent:
        p.paragraph_format.first_line_indent = Pt(21)
    run = p.add_run(text)
    _set_font(run, size, color=color)
    return p


def _kv_table(doc: Document, rows: list[tuple[str, str]]):
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.columns[0].width = Cm(3.6)
    for i, (k, v) in enumerate(rows):
        for j, (text, bold) in enumerate(((k, True), (v, False))):
            cell = table.cell(i, j)
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(text)
            _set_font(run, 10, bold=bold, color=_GRAY if j == 0 else _PRIMARY)
    doc.add_paragraph()


def build_script_docx(bundle: dict[str, Any]) -> BytesIO:
    """脚本套装 -> Word。bundle 与 script_service 返回结构一致。"""
    doc = Document()
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Cm(2.4)
        section.left_margin = section.right_margin = Cm(2.6)

    # 封面标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("短视频脚本策划案")
    _set_font(r, 20, bold=True)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub.add_run(bundle.get("topic", ""))
    _set_font(r2, 12, color=_GRAY)
    doc.add_paragraph()

    _kv_table(doc, [
        ("生成主题", bundle.get("topic", "")),
        ("目标平台", bundle.get("platform_name", bundle.get("platform", ""))),
        ("视频时长", f"{bundle.get('duration', '')} 秒"),
        ("内容风格", bundle.get("style", "")),
        ("生成模型", bundle.get("source_model", "")),
    ])

    _heading(doc, "一、主题概述")
    _body(doc, bundle.get("overview", ""))

    _heading(doc, "二、开场爆款钩子（3 秒）")
    _body(doc, bundle.get("hook", ""))

    _heading(doc, "三、分镜脚本")
    segments = bundle.get("segments", [])
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    header = ["时间", "类型", "画面内容", "出镜台词", "字幕重点"]
    for j, h in enumerate(header):
        cell = table.rows[0].cells[j]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        _set_font(run, 9.5, bold=True, color=_BLUE)
    for seg in segments:
        row = table.add_row()
        vals = [
            f"{seg.get('start_time', '')}-{seg.get('end_time', '')}",
            "出镜" if seg.get("type") == "on_camera" else "旁白",
            seg.get("scene", ""), seg.get("lines", ""), seg.get("subtitle", ""),
        ]
        for j, v in enumerate(vals):
            cell = row.cells[j]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(v))
            _set_font(run, 9.5)
    doc.add_paragraph()

    _heading(doc, "四、结尾互动引导")
    _body(doc, bundle.get("ending", ""))

    _heading(doc, "五、爆款标题（10 组）")
    title_items = bundle.get("title_items") or [
        {"title": t, "tone": "标题"} for t in bundle.get("titles", [])
    ]
    for i, item in enumerate(title_items, 1):
        _body(doc, f"{i}. {item.get('title', item)}（{item.get('tone', '')}）", size=10.5)

    _heading(doc, "六、话题标签矩阵")
    tiers = {"hot": [], "mid": [], "long": []}
    for tag in bundle.get("tags", []):
        tiers.setdefault(tag.get("tier", "mid"), []).append(tag.get("text", ""))
    if not any(tiers.values()):
        tiers = {k: v for k, v in (bundle.get("tags") or {}).items()}  # 兼容 dict 结构
    for tier, label in _TIER_NAMES.items():
        items = tiers.get(tier, [])
        _body(doc, f"{label}（{len(items)}）：{' '.join(items) if isinstance(items, list) else items}", size=10.5)

    _heading(doc, "七、TTS 配音文稿（可直接用于剪映配音）")
    tts = bundle.get("tts_text", "")
    for line in (bundle.get("tts_sentences") or [s for s in tts.split("\n") if s]):
        _body(doc, line, size=10.5)

    if bundle.get("body_text"):
        _heading(doc, "八、完整正文")
        for line in bundle.get("body_text", "").split("\n"):
            if line.strip():
                _body(doc, line.strip(), size=10.5)

    _footer(doc)
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def build_simple_docx(title: str, sections: list[dict[str, Any]]) -> BytesIO:
    """通用导出：sections = [{heading, body | lines[] | rows:[[..]]}]。"""
    doc = Document()
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Cm(2.4)
        section.left_margin = section.right_margin = Cm(2.6)
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_font(t.add_run(title), 18, bold=True)
    doc.add_paragraph()

    for idx, sec in enumerate(sections, 1):
        _heading(doc, f"{idx}. {sec['heading']}")
        for line in sec.get("lines", []):
            _body(doc, line)
        if sec.get("rows"):
            table = doc.add_table(rows=len(sec["rows"]), cols=len(sec["rows"][0]))
            table.style = "Table Grid"
            for i, row in enumerate(sec["rows"]):
                for j, v in enumerate(row):
                    cell = table.cell(i, j)
                    cell.text = ""
                    run = cell.paragraphs[0].add_run(str(v))
                    _set_font(run, 10, bold=(i == 0), color=_BLUE if i == 0 else _PRIMARY)
            doc.add_paragraph()
    _footer(doc)
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _footer(doc: Document):
    from docx.oxml import OxmlElement

    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run._element.addnext(fld)


# ---------------------------------------------------------------- 记录 -> 文档
def record_to_script_bundle(r) -> dict[str, Any]:
    """ContentRecord(script) -> build_script_docx 所需的套装结构。"""
    from app.services.prompts.templates import PLATFORM_PROFILES

    content = r.content or {}
    return {
        "topic": r.topic, "platform": r.platform,
        "platform_name": PLATFORM_PROFILES.get(r.platform, {}).get("label", r.platform),
        "duration": r.duration, "style": r.style,
        "overview": content.get("topic_overview", ""),
        "hook": content.get("hook", ""),
        "segments": content.get("segments", []),
        "ending": content.get("ending", ""),
        "titles": r.titles or [],
        "title_items": content.get("titles", []),
        "tags": r.tags or [],
        "tts_text": r.tts_text or "",
        "tts_sentences": [s for s in (r.tts_text or "").split("\n") if s],
        "body_text": r.body_text or "",
        "source_model": r.source_model,
    }


def record_to_docx(r) -> BytesIO | None:
    """按记录类型生成 Word；返回 None 表示该类型不支持。"""
    if r.record_type == "script":
        return build_script_docx(record_to_script_bundle(r))
    if r.record_type == "titles":
        sections = [
            {"heading": "爆款标题", "lines": [f"{i}. {t}" for i, t in enumerate(r.titles or [], 1)]},
            {"heading": "话题标签", "lines": [f"{t.get('text', '')}" for t in (r.tags or [])]},
        ]
        return build_simple_docx(f"标题&标签 - {r.topic}", sections)
    if r.record_type == "copywriting":
        sections = [{"heading": "改写成稿", "lines": [ln for ln in (r.body_text or "").split("\n") if ln.strip()]}]
        return build_simple_docx(f"文案{' - ' + r.style if r.style else ''}", sections)
    return None


def build_records_zip(records: list) -> BytesIO:
    """多记录打包导出（批量结果 / 历史批量导出）：records.zip，文件名按主题命名。"""
    import io
    import re
    import zipfile

    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for r in records:
            docx = record_to_docx(r)
            if docx is None:
                continue
            safe = re.sub(r'[\\/:*?"<>|]', "_", r.topic or "untitled")[:40]
            zf.writestr(f"{r.record_type}_{safe}.docx", docx.getvalue())
    buf.seek(0)
    return buf
